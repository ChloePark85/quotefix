import streamlit as st
import pandas as pd
from docx import Document
import io

def normalize_quotes(text):
    # 모든 따옴표를 곡선 따옴표로 변환
    import re
    
    # 모든 따옴표를 임시 플레이스홀더로 변경
    text = re.sub(r'["\u201C\u201D]', '§DQUOTE§', text)
    text = re.sub(r"['\u2018\u2019]", '§SQUOTE§', text)
    
    # 큰따옴표 처리 (홀수번째=열기, 짝수번째=닫기)
    parts = text.split('§DQUOTE§')
    result_parts = []
    for i, part in enumerate(parts):
        result_parts.append(part)
        if i < len(parts) - 1:
            if i % 2 == 0:
                result_parts.append('\u201C')  # " (열기)
            else:
                result_parts.append('\u201D')  # " (닫기)
    text = ''.join(result_parts)
    
    # 작은따옴표 처리 (홀수번째=열기, 짝수번째=닫기)
    parts = text.split('§SQUOTE§')
    result_parts = []
    for i, part in enumerate(parts):
        result_parts.append(part)
        if i < len(parts) - 1:
            if i % 2 == 0:
                result_parts.append('\u2018')  # ' (열기)
            else:
                result_parts.append('\u2019')  # ' (닫기)
    text = ''.join(result_parts)
    
    return text

def count_quotes(text):
    # 다양한 따옴표 타입 카운트
    straight_double = text.count('"')
    straight_single = text.count("'")
    curly_double = text.count('\u201C') + text.count('\u201D')  # " "
    curly_single = text.count('\u2018') + text.count('\u2019')  # ' '
    
    quote_types = {
        '직선 큰따옴표 (")': straight_double,
        "직선 작은따옴표 (')": straight_single,
        '곡선 큰따옴표': curly_double,
        "곡선 작은따옴표": curly_single,
    }
    return quote_types

def main():
    st.set_page_config(
        page_title="따옴표 통일 도구",
        page_icon="📝",
        layout="wide"
    )
    
    st.title("📝 원고 따옴표 통일 도구")
    st.markdown("원고를 업로드하면 모든 따옴표를 **곡선 따옴표(" ")로 자동 통일**합니다.")

    # 메인 컨텐츠
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📥 원본 파일 업로드")
        
        st.info("💡 **파일을 드래그해서 놓으세요**\n"
               "지원 형식: TXT, DOCX")
        
        upload_file = st.file_uploader(
            "파일을 여기로 드래그해서 놓으세요",
            type=['txt', 'docx'],
            label_visibility="collapsed"
        )
        
        original_text = ""
        
        if upload_file is not None:
            try:
                if upload_file.name.endswith('.txt'):
                    bytes_data = upload_file.read()
                    try:
                        original_text = bytes_data.decode('utf-8')
                    except UnicodeDecodeError:
                        original_text = bytes_data.decode('cp949', errors='ignore')
                
                elif upload_file.name.endswith('.docx'):
                    doc = Document(io.BytesIO(upload_file.read()))
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    original_text = '\n\n'.join(paragraphs)
                
                if original_text.strip():
                    st.success(f"✅ 파일 업로드 완료: {upload_file.name}")
                else:
                    st.warning("⚠️ 파일에서 텍스트를 찾을 수 없습니다.")
                    
            except Exception as e:
                st.error(f"❌ 파일 읽기 오류: {e}")
                original_text = ""
        
        if original_text:
            st.info(f"텍스트 길이: {len(original_text):,}자")
            
            with st.expander("📈 따옴표 현황"):
                quote_counts = count_quotes(original_text)
                for quote_type, count in quote_counts.items():
                    if count > 0:
                        st.write(f"- {quote_type}: {count}개")
    
    with col2:
        st.subheader("📤 곡선 따옴표로 변환")
        
        if st.button("🔄 곡선 따옴표로 변환하기", type="primary", use_container_width=True):
            if original_text.strip():
                with st.spinner("곡선 따옴표로 변환 중..."):
                    normalized_text = normalize_quotes(original_text)
                
                st.text_area(
                    "변환 결과",
                    value=normalized_text,
                    height=300,
                    key="result"
                )
                
                if original_text != normalized_text:
                    st.success("✅ 곡선 따옴표로 변환 완료!")
                    
                    # Word 다운로드만 제공
                    try:
                        from docx.shared import Inches
                        
                        # 새 워드 문서 생성
                        doc = Document()
                        
                        # 문서 여백 설정
                        sections = doc.sections
                        for section in sections:
                            section.top_margin = Inches(1)
                            section.bottom_margin = Inches(1)
                            section.left_margin = Inches(1)
                            section.right_margin = Inches(1)
                        
                        # 텍스트를 문단별로 나누어 추가
                        paragraphs = normalized_text.split('\n')
                        for para_text in paragraphs:
                            if para_text.strip():  # 빈 줄이 아닌 경우
                                doc.add_paragraph(para_text)
                            else:  # 빈 줄인 경우
                                doc.add_paragraph("")
                        
                        # 메모리에서 워드 파일 생성
                        doc_buffer = io.BytesIO()
                        doc.save(doc_buffer)
                        doc_buffer.seek(0)
                        
                        st.download_button(
                            "📝 Word 문서 다운로드",
                            data=doc_buffer.getvalue(),
                            file_name="곡선따옴표_통일.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error(f"Word 파일 생성 오류: {e}")
                else:
                    st.info("변환할 따옴표가 없습니다.")
            else:
                st.warning("파일을 업로드해주세요.")
        else:
            st.info("파일을 업로드하고 변환 버튼을 클릭하세요.")

if __name__ == "__main__":
    main()